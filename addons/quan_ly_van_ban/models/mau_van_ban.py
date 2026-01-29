# -*- coding: utf-8 -*-
from odoo import models, fields, api
from odoo.exceptions import ValidationError
import base64
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
import re
from datetime import datetime

class MauVanBan(models.Model):
    _name = 'mau_van_ban'
    _description = 'Mẫu văn bản'
    _inherit = ['mail.thread', 'mail.activity.mixin']
    _order = 'sequence, name'
    
    name = fields.Char("Tên mẫu", required=True)
    ma_mau = fields.Char("Mã mẫu", required=True, copy=False, index=True)
    sequence = fields.Integer("Thứ tự", default=10)
    loai_van_ban_id = fields.Many2one('loai_van_ban', string="Loại văn bản", required=True)
    
    # Nội dung mẫu
    noi_dung_html = fields.Html("Nội dung mẫu (HTML)", 
                                 help="Sử dụng biến: ${ten_nhan_vien}, ${ten_khach_hang}, ${ngay_hom_nay}, ${so_van_ban}")
    
    # File Word template (nếu có)
    file_word_template = fields.Binary("File Word mẫu", attachment=True)
    file_word_template_name = fields.Char("Tên file Word")
    
    # Hướng dẫn sử dụng
    huong_dan = fields.Text("Hướng dẫn sử dụng")
    
    # Danh sách biến có thể dùng
    bien_kha_dung = fields.Text("Biến khả dụng", compute="_compute_bien_kha_dung", store=False,
                                 help="Danh sách các biến có thể sử dụng trong mẫu")
    
    # Trạng thái
    trang_thai = fields.Selection([
        ('draft', 'Nháp'),
        ('active', 'Đang sử dụng'),
        ('archived', 'Lưu trữ')
    ], string="Trạng thái", default='draft', required=True)
    
    active = fields.Boolean("Active", default=True)
    
    # Thống kê
    so_lan_su_dung = fields.Integer("Số lần sử dụng", default=0, readonly=True)
    
    _sql_constraints = [
        ('ma_mau_unique', 'unique(ma_mau)', 'Mã mẫu phải là duy nhất!'),
    ]
    
    @api.depends('loai_van_ban_id')
    def _compute_bien_kha_dung(self):
        """Tính toán danh sách biến khả dụng dựa trên loại văn bản"""
        for record in self:
            bien_chung = [
                "${so_van_ban} - Số văn bản",
                "${ngay_hom_nay} - Ngày hôm nay",
                "${ngay_thang_nam} - Ngày tháng năm đầy đủ",
                "${nam_hien_tai} - Năm hiện tại"
            ]
            
            bien_nhan_vien = [
                "${ten_nhan_vien} - Tên nhân viên",
                "${ma_nhan_vien} - Mã nhân viên",
                "${chuc_vu} - Chức vụ",
                "${phong_ban} - Phòng ban",
                "${email_nhan_vien} - Email nhân viên",
                "${sdt_nhan_vien} - Số điện thoại nhân viên"
            ]
            
            bien_khach_hang = [
                "${ten_khach_hang} - Tên khách hàng",
                "${ma_khach_hang} - Mã khách hàng",
                "${email_khach_hang} - Email khách hàng",
                "${sdt_khach_hang} - Số điện thoại khách hàng",
                "${dia_chi_khach_hang} - Địa chỉ khách hàng"
            ]
            
            bien_van_ban = [
                "${trich_yeu} - Trích yếu",
                "${noi_dung} - Nội dung văn bản"
            ]
            
            all_vars = bien_chung + bien_nhan_vien + bien_khach_hang + bien_van_ban
            record.bien_kha_dung = "\n".join(all_vars)
    
    def action_set_active(self):
        """Đưa mẫu vào trạng thái đang sử dụng"""
        self.write({'trang_thai': 'active'})
    
    def action_set_draft(self):
        """Đưa mẫu về trạng thái nháp"""
        self.write({'trang_thai': 'draft'})
    
    def action_archive(self):
        """Lưu trữ mẫu"""
        self.write({'trang_thai': 'archived', 'active': False})
    
    def action_tao_van_ban_tu_mau(self):
        """Mở wizard tạo văn bản từ mẫu"""
        self.ensure_one()
        return {
            'name': 'Tạo văn bản từ mẫu',
            'type': 'ir.actions.act_window',
            'res_model': 'wizard.tao.van.ban.tu.mau',
            'view_mode': 'form',
            'target': 'new',
            'context': {
                'default_mau_van_ban_id': self.id,
                'default_loai_van_ban_id': self.loai_van_ban_id.id,
            }
        }
    
    def thay_the_bien(self, noi_dung, data_dict):
        """
        Thay thế các biến trong nội dung bằng dữ liệu thực
        
        :param noi_dung: Nội dung chứa biến ${...}
        :param data_dict: Dictionary chứa giá trị thực
        :return: Nội dung đã thay thế
        """
        if not noi_dung:
            return ""
        
        result = noi_dung
        
        # Thay thế các biến
        for key, value in data_dict.items():
            pattern = r'\$\{' + key + r'\}'
            result = re.sub(pattern, str(value or ''), result)
        
        return result
    
    def tao_du_lieu_mac_dinh(self):
        """Tạo dữ liệu mặc định cho các biến"""
        today = fields.Date.today()
        
        data = {
            'ngay_hom_nay': today.strftime('%d/%m/%Y'),
            'ngay_thang_nam': f"Ngày {today.day} tháng {today.month} năm {today.year}",
            'nam_hien_tai': str(today.year),
            'so_van_ban': 'NEW',
        }
        
        return data
    
    def xuat_word(self, data_dict=None):
        """
        Xuất file Word từ mẫu với dữ liệu được merge
        
        :param data_dict: Dictionary chứa dữ liệu để merge
        :return: Binary data của file Word
        """
        self.ensure_one()
        
        if not data_dict:
            data_dict = self.tao_du_lieu_mac_dinh()
        
        # Tạo document mới
        doc = Document()
        
        # Thiết lập font và style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        
        # Thêm tiêu đề
        heading = doc.add_heading(self.name, level=1)
        heading.alignment = 1  # Center
        
        # Thay thế biến trong nội dung HTML và thêm vào document
        if self.noi_dung_html:
            # Convert HTML sang plain text (đơn giản)
            noi_dung = self.noi_dung_html
            # Loại bỏ HTML tags
            noi_dung = re.sub(r'<[^>]+>', '', noi_dung)
            # Thay thế biến
            noi_dung = self.thay_the_bien(noi_dung, data_dict)
            
            # Thêm các đoạn văn
            for paragraph in noi_dung.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        
        # Lưu vào BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Tăng số lần sử dụng
        self.sudo().write({'so_lan_su_dung': self.so_lan_su_dung + 1})
        
        return file_stream.read()
